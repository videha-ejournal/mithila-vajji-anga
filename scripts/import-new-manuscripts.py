from __future__ import annotations

import json
import re
import time
import html
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
APP = ROOT / "app"
IDEAS_SOURCE = Path(r"C:\Users\DELL\Downloads\२५.docx")
IDEAS_CHAPTER_46 = Path(r"C:\Users\DELL\Downloads\४६.docx")
POLITICAL_HISTORY_SOURCE = Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\KDP PAPERBACK\HISTORY_MITHILA_ANGA_VAJJI_KDP_PAGE636_FINAL_SAFE.pdf")
HISTORY_SOURCES = [
    Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\00_Volume_2\Part 1-3\History_Mithila_Vajji_Anga_Volume_II_Cumulative_Chapter25_6x9_Hardback.docx"),
    Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\00_Volume_2\Part 4\History_Mithila_Vajji_Anga_Volume_II_Cumulative_Part_26_onward_Chapter35_6x9_Hardback.docx"),
    Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\00_Volume_2\Part 5\History_Mithila_Vajji_Anga_Volume_II_Cumulative_Part_36_onward_Chapter43_6x9_Hardback.docx"),
    Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\00_Volume_2\Part 6\History_Mithila_Vajji_Anga_Volume_II_Cumulative_Part_44_onward_Chapter51_6x9_Hardback.docx"),
    Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\00_Volume_2\Part 7\History_Mithila_Vajji_Anga_Volume_II_Cumulative_Part_52_onward_Chapter64_6x9_Hardback.docx"),
    Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\00_Volume_2\Part 8\History_Mithila_Vajji_Anga_Volume_II_Cumulative_Part_65_onward_Chapter73_6x9_Hardback.docx"),
    Path(r"C:\Users\DELL\Desktop\01_DONE_FINAL\FINAL\History_of_Mithila_Vajji_Anga\00_Volume_2\Part 9\History_Mithila_Vajji_Anga_Volume_II_Cumulative_Part_74_onward_Chapter86_6x9_Hardback.docx"),
]

TRANSLATION_CACHE = ROOT / "work" / "maithili-english-cache.json"
SECTION_LABELS = [
    "समस्या", "मूल प्रतिपादन", "प्रमुख तर्क", "पूर्वपक्ष", "उत्तरपक्ष",
    "भारतीय संवाद", "मिथिलाक समानान्तर दृष्टि", "समकालीन प्रयोग",
    "अध्याय-निष्कर्ष", "अध्याय-ग्रन्थसूची",
]
SECTION_NAMES = {
    "समस्या": "Problem and scope",
    "मूल प्रतिपादन": "Central thesis",
    "प्रमुख तर्क": "Major arguments",
    "पूर्वपक्ष": "Pūrvapakṣa",
    "उत्तरपक्ष": "Uttarapakṣa",
    "भारतीय संवाद": "Indian philosophical dialogue",
    "मिथिलाक समानान्तर दृष्टि": "Mithila’s parallel perspective",
    "समकालीन प्रयोग": "Contemporary applications",
    "अध्याय-निष्कर्ष": "Chapter conclusion",
    "अध्याय-ग्रन्थसूची": "Chapter bibliography",
}

ENGLISH_TITLES = [
    "The Second Turn in Philosophy: From Consciousness to Language",
    "Why Philosophy Changed in the Twentieth Century",
    "The Division between Analytic and Continental Philosophy",
    "Where Does Indian Philosophy Sit on This Map?",
    "Mithila’s Parallel Point of Entry",
    "The Risks of Comparative Philosophy",
    "Frege: Logic, Meaning and Reference",
    "Russell: Description, Logical Form and the World",
    "Moore and Common Sense",
    "Early Wittgenstein: Language, Fact, Picture and the Limits of Expression",
    "Logical Positivism: Verification, Science and the Limits of Metaphysics",
    "Carnap and the Question of Metaphysics",
    "Quine: The Crisis of the Analytic–Synthetic Distinction",
    "Analytic Philosophy and Navya-Nyāya",
    "Later Wittgenstein: Meaning in Use",
    "Austin: Saying Is Also Doing",
    "Searle and Speech Acts",
    "Can Language Be Private?",
    "The Meaning of a Word: Object, Use or Relation?",
    "Indian Theories of Verbal Meaning",
    "Mithila’s Tradition of Linguistic Analysis",
    "What Is Structure?",
    "Saussure: Signifier and Signified",
    "Lévi-Strauss: Myth and Structure",
    "Structure and the Individual",
    "Indian Linguistic Traditions and Structure",
    "The Limits of Structuralism",
    "From Structuralism to Post-Structuralism",
    "Derrida: A Critique of the Metaphysics of Presence",
    "Différance: Deferral and Difference in Meaning",
    "Writing, Speech and the Question of the Centre",
    "What Deconstruction Is and Is Not",
    "Indian Philosophy and Deconstruction",
    "Navya-Nyāya as a Counterpoint to Deconstruction",
    "A Parallel Uttarapakṣa",
    "Can Knowledge and Power Be Separated?",
    "Foucault’s Archaeology",
    "Genealogy: Another Way of Writing History",
    "Discipline, Surveillance and Normalisation",
    "Body, Institution and Knowledge",
    "Foucault and the Parallel View of History",
    "What Is Postmodernism?",
    "Lyotard and Incredulity toward Metanarratives",
    "Baudrillard: Simulacra, Media and Hyperreality",
    "Has Truth Ended?",
    "Science, History and Postmodern Critique",
    "Parallel Philosophy’s Reply",
    "Hermeneutics: Why Is Understanding Itself a Problem?",
    "Gadamer: Tradition and Pre-understanding",
    "Ricoeur: Text, Symbol and Interpretation",
    "Authorial Intention and the Reader’s Meaning",
    "The Indian Commentary Tradition",
    "Mithila’s Śāstrārtha as a Hermeneutic Practice?",
    "The Frankfurt School",
    "Horkheimer and Adorno: The Self-Critique of Modern Reason",
    "Marcuse: Consumer Society and One-Dimensional Man",
    "Habermas: Communicative Reason",
    "The Public Sphere and Democracy",
    "Śāstrārtha, Debate and Communicative Reason",
    "The Many Voices of Feminist Philosophy",
    "Simone de Beauvoir: Woman Is Made",
    "Body, Experience and Power",
    "Judith Butler: Gender, Performativity and Identity",
    "Intersectionality: Caste, Class, Gender, Religion and Region",
    "Questions for Indian Feminist Philosophy",
    "Mithila’s Parallel Feminist Perspective",
    "What Does Colonial Knowledge Do?",
    "Edward Said: Orientalism",
    "Gayatri Spivak: Can the Subaltern Speak?",
    "Homi Bhabha: Hybridity and the Colonial Interstice",
    "Subaltern Studies and the Margins of History",
    "From Postcolonial to Decolonial Thought",
    "The Colonial Classification of Indian Philosophy",
    "Mithila between Empire, Nation and Region",
    "New Questions in the Philosophy of Mind",
    "The Mind–Body Problem",
    "Behaviourism, Identity Theory and Functionalism",
    "The Hard Problem of Consciousness",
    "Intentionality and the Self",
    "Indian Debates on Self, No-Self and Mind",
    "A Three-Way Dialogue: Nyāya, Buddhism and Modern Philosophy of Mind",
    "New Debates in Scientific Realism",
    "Natural Science and Social Science",
    "Technology Is More Than a Tool",
    "Truth in the Digital World",
    "Algorithms, Classification and Power",
    "Artificial Intelligence: What Is Intelligence?",
    "Machine Consciousness: An Open Question",
    "Evidence and Responsibility in the Age of AI",
    "A Second Form of Multi-Source Rationalism",
    "Coexistential Realism after the Linguistic Turn",
    "Dignity-Based Ethics and New Identities",
    "Polycentric Democracy and Network Power",
    "Rational Religious Pluralism in a Post-Secular Society",
    "The Parallel View of History: Genealogy, Archive and Evidence",
    "Self-Critique: How Can Parallel Philosophy Accept Its Own Deconstruction?",
    "Structure Exists, but Structure Is Not Destiny",
    "Language Is a Limit, but Not a Prison",
    "Power Influences Knowledge, but Does Not Make Truth Arbitrary",
    "Parallel Philosophy: Reconstruction after Deconstruction",
]

ENGLISH_SYNOPSIS = [
    "The chapter traces philosophy’s movement from the isolated knowing subject toward language, meaning, logical form and shared rules.",
    "The chapter examines how science, symbolic logic, war, colonialism and the modern administrative state forced philosophy to reconsider its subject and method.",
    "The chapter tests the institutional divide between analytic and continental philosophy instead of treating it as a complete map of twentieth-century thought.",
    "The chapter criticises the separation of Indian philosophy from general philosophical history and restores its debates about knowledge, language, causation and reality to a shared field of inquiry.",
    "The chapter places Mithila’s Nyāya, Navya-Nyāya, Mīmāṃsā, commentary and debate traditions in conversation with modern questions without claiming premature equivalence.",
    "The chapter distinguishes comparison from cultural triumphalism: similar questions, terms or conclusions do not by themselves prove identical theories or historical influence.",
    "The chapter examines Frege’s formal logic and his distinction between sense and reference alongside Indian debates about verbal meaning.",
    "The chapter explains Russell’s theory of descriptions and the difference between grammatical appearance and logical form.",
    "The chapter studies Moore’s defence of common sense and compares its realist commitments with questions of perception and warrant in Nyāya.",
    "The chapter reads early Wittgenstein through facts, pictures, logical form and the limits of what language can state.",
    "The chapter assesses verificationism, its critique of metaphysics and the problems created by its own criterion of meaning.",
    "The chapter asks whether philosophical problems are merely products of linguistic confusion or retain genuine metaphysical force.",
    "The chapter examines Quine’s challenge to the analytic–synthetic distinction and the image of knowledge as an interconnected web.",
    "The chapter compares analytic philosophy and Navya-Nyāya through technical language, relations and inference while preserving their historical differences.",
    "The chapter presents meaning as use through language-games, rule-following and forms of life.",
    "The chapter examines Austin’s claim that utterances can perform social acts when conventions, authority and circumstances are in place.",
    "The chapter develops Searle’s account of speech acts through intention, convention, institutional rules and indirect requests.",
    "The chapter asks whether a language whose standards depend entirely on private sensation can sustain a distinction between correct use and merely seeming correct.",
    "The chapter compares object, use and relational accounts of word meaning and tests each against difficult cases such as justice, pain and number.",
    "The chapter introduces Nyāya, Mīmāṃsā, grammatical and Buddhist approaches to word meaning, sentence cognition, intention and testimony.",
    "The chapter identifies Mithila’s practices of definition, qualification, burden of proof and objection as a disciplined tradition of linguistic analysis.",
    "The chapter defines structure as a system of relations, differences, rules and transformations rather than a mere collection of separate elements.",
    "The chapter explains Saussure’s signifier and signified as two sides of the linguistic sign operating within a system of differences.",
    "The chapter examines Lévi-Strauss’s structural reading of myth through recurring relations, oppositions and transformations across variant narratives.",
    "The chapter asks how agency and moral responsibility remain possible when language, institutions and classifications provide structures prior to individual action.",
]

PARTS = [
    (1, 6, "Twentieth-Century Philosophical Turns"),
    (7, 14, "Analytic Philosophy and the New World of Logic"),
    (15, 21, "The Linguistic Turn and the Philosophy of Meaning"),
    (22, 27, "Structuralism"),
    (28, 35, "Post-Structuralism and Deconstruction"),
    (36, 41, "Foucault, Knowledge and Power"),
    (42, 47, "Postmodernity"),
    (48, 53, "Interpretation, Text and Reader"),
    (54, 59, "Critical Theory after Marx"),
    (60, 66, "Feminism, Gender and the Body"),
    (67, 74, "Colonialism, Postcolonialism and Knowledge"),
    (75, 81, "Mind, Consciousness and Personhood"),
    (82, 89, "Science, Technology, the Digital World and AI"),
    (90, 96, "Parallel Philosophy in the Twenty-First Century"),
    (97, 100, "New Sūtras"),
]


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def part_for(number: int) -> str:
    return next(label for start, end, label in PARTS if start <= number <= end)


def load_translation_cache() -> dict[str, str]:
    if TRANSLATION_CACHE.exists():
        return json.loads(TRANSLATION_CACHE.read_text(encoding="utf-8"))
    return {}


def translate_maithili(text: str, cache: dict[str, str]) -> str:
    text = clean(text)
    if not text:
        return ""
    if text in cache:
        return cache[text]
    payload = urllib.parse.urlencode({
        "client": "gtx", "sl": "mai", "tl": "en", "dt": "t", "q": text,
    }).encode("utf-8")
    request = urllib.request.Request(
        "https://translate.googleapis.com/translate_a/single",
        data=payload,
        headers={"User-Agent": "Mozilla/5.0", "Content-Type": "application/x-www-form-urlencoded;charset=UTF-8"},
    )
    last_error = None
    for attempt in range(2):
        try:
            with urllib.request.urlopen(request, timeout=30) as response:
                data = json.loads(response.read().decode("utf-8"))
            translated = clean("".join(piece[0] for piece in data[0] if piece and piece[0]))
            cache[text] = translated
            TRANSLATION_CACHE.parent.mkdir(parents=True, exist_ok=True)
            TRANSLATION_CACHE.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")
            time.sleep(0.35)
            return translated
        except Exception as error:
            last_error = error
            time.sleep(2 * (attempt + 1))

    chunks: list[str] = []
    remaining = text
    while remaining:
        if len(remaining) <= 1000:
            chunks.append(remaining)
            break
        split_at = max(remaining.rfind(mark, 0, 1000) for mark in ["।", ".", " "])
        split_at = split_at + 1 if split_at > 500 else 1000
        chunks.append(remaining[:split_at].strip())
        remaining = remaining[split_at:].strip()
    translated_chunks = []
    for chunk in chunks:
        mobile_query = urllib.parse.urlencode({"sl": "mai", "tl": "en", "q": chunk})
        mobile_request = urllib.request.Request(
            f"https://translate.google.com/m?{mobile_query}",
            headers={"User-Agent": "Mozilla/5.0 (Linux; Android 10)"},
        )
        for attempt in range(5):
            try:
                with urllib.request.urlopen(mobile_request, timeout=30) as response:
                    page = response.read().decode("utf-8")
                match = re.search(r'class="result-container">(.*?)</div>', page, re.DOTALL)
                if not match:
                    raise RuntimeError("Google mobile translation result was absent")
                translated_chunks.append(clean(html.unescape(re.sub(r"<[^>]+>", "", match.group(1)))))
                time.sleep(0.25)
                break
            except Exception as error:
                last_error = error
                time.sleep(3 * (attempt + 1))
        else:
            raise RuntimeError(f"Translation failed after API and mobile retries: {last_error}")
    translated = clean(" ".join(translated_chunks))
    cache[text] = translated
    TRANSLATION_CACHE.parent.mkdir(parents=True, exist_ok=True)
    TRANSLATION_CACHE.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")
    return translated


def clipped_join(paragraphs: list[str], limit: int) -> str:
    result = ""
    for paragraph in paragraphs:
        candidate = clean(f"{result} {paragraph}")
        if len(candidate) > limit:
            if not result:
                result = candidate[:limit].rsplit(" ", 1)[0]
            break
        result = candidate
    return result


def parse_maithili_idea_chapters() -> dict[int, dict[str, object]]:
    document = Document(IDEAS_CHAPTER_46)
    starts: list[tuple[int, int, str]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = clean(paragraph.text)
        match = re.match(r"^अध्याय\s+(\d+)\s*[—-]\s*(.+)$", text)
        if match and 1 <= int(match.group(1)) <= 46:
            starts.append((index, int(match.group(1)), match.group(2)))
    parsed: dict[int, dict[str, object]] = {}
    for position, (start, number, title) in enumerate(starts):
        end = starts[position + 1][0] if position + 1 < len(starts) else len(document.paragraphs)
        buckets: dict[str, list[str]] = {label: [] for label in SECTION_LABELS}
        active = "समस्या"
        for paragraph in document.paragraphs[start + 1:end]:
            text = clean(paragraph.text)
            if not text:
                continue
            normalized_label = re.sub(r"^\d+\.\s*", "", text)
            if normalized_label in SECTION_LABELS:
                active = normalized_label
                continue
            if text.startswith("चित्र ") or text.startswith("तालिका "):
                continue
            buckets[active].append(text)
        parsed[number] = {"title": title, "buckets": buckets}
    if sorted(parsed) != list(range(1, 47)):
        raise RuntimeError(f"Expected supplied ideas 1–46; found {sorted(parsed)}")
    return parsed


def build_ideas():
    ideas_path = APP / "ideas-volume2.json"
    ideas = json.loads(ideas_path.read_text(encoding="utf-8"))
    source_chapters = parse_maithili_idea_chapters()
    cache = load_translation_cache()
    for record in ideas:
        number = record["number"]
        if number > 46:
            record["status"] = "Planned"
            record["summary"] = "This chapter remains an approved plan entry. No completed English argument is claimed."
            record["source"] = "Gajendra Thakur’s Parallel Philosophy, Volume II · approved plan entry"
            record.pop("purvapaksha", None)
            record.pop("uttarapaksha", None)
            record.pop("synthesis", None)
            record.pop("sections", None)
            continue
        chapter = source_chapters[number]
        buckets = chapter["buckets"]
        summary_source = clipped_join(buckets["समस्या"][:3] + buckets["मूल प्रतिपादन"][:3], 4300)
        purva_source = clipped_join(buckets["पूर्वपक्ष"][:5], 3600)
        uttara_source = clipped_join(buckets["उत्तरपक्ष"][:5], 3600)
        synthesis_source = clipped_join(buckets["अध्याय-निष्कर्ष"][:4], 3600)
        section_index = []
        for label in SECTION_LABELS:
            count = len(buckets[label])
            if count:
                suffix = f" · {count} source passage{'s' if count != 1 else ''}"
                section_index.append(f"{SECTION_NAMES[label]}{suffix}")
        record.update({
            "title": ENGLISH_TITLES[number - 1],
            "part": part_for(number),
            "status": "Available in English",
            "summary": translate_maithili(summary_source, cache),
            "purvapaksha": translate_maithili(purva_source, cache),
            "uttarapaksha": translate_maithili(uttara_source, cache),
            "synthesis": translate_maithili(synthesis_source, cache),
            "sections": section_index,
            "source": f"Gajendra Thakur’s Parallel Philosophy, Volume II · supplied Maithili Chapter {number} · English translation",
        })
    ideas_path.write_text(json.dumps(ideas, ensure_ascii=False, indent=2), encoding="utf-8")
    return ideas


def enrich_history():
    research_path = APP / "research-data.json"
    research = json.loads(research_path.read_text(encoding="utf-8"))
    imported = []
    for path in HISTORY_SOURCES:
        document = Document(path)
        paragraphs = document.paragraphs
        starts = []
        for index, paragraph in enumerate(paragraphs):
            text = clean(paragraph.text)
            match = re.match(r"^Chapter\s+(\d+)\s*[—-]\s*(.+)$", text, re.IGNORECASE)
            if match and 1 <= int(match.group(1)) <= 86:
                starts.append((index, int(match.group(1)), match.group(2)))
        for position, (start, number, title) in enumerate(starts):
            end = starts[position + 1][0] if position + 1 < len(starts) else len(paragraphs)
            sections = []
            body = []
            for paragraph in paragraphs[start + 1 : end]:
                text = clean(paragraph.text)
                if not text:
                    continue
                if paragraph.style and paragraph.style.name == "Heading 2" and re.match(rf"^{number}\.\d+", text):
                    sections.append(text)
                elif len(text) >= 120 and "Bibliography" not in text and len(body) < 4:
                    body.append(text)
            record = next((chapter for chapter in research["social"] if chapter["number"] == number), None)
            if not record:
                continue
            record["title"] = title
            record["status"] = "Complete"
            record["pages"] = f"Supplied cumulative manuscript · Chapters {starts[0][1]}–{starts[-1][1]}"
            record["summary"] = " ".join(body)[:2400]
            record["sections"] = sections
            imported.append({"number": number, "title": title, "sections": len(sections), "source": path.name})
    research_path.write_text(json.dumps(research, ensure_ascii=False, indent=2), encoding="utf-8")
    (APP / "supplied-history-additions.json").write_text(json.dumps(imported, ensure_ascii=False, indent=2), encoding="utf-8")
    return imported


def enrich_political_history_sections():
    research_path = APP / "research-data.json"
    research = json.loads(research_path.read_text(encoding="utf-8"))
    reader = PdfReader(POLITICAL_HISTORY_SOURCE)
    outline = reader.outline
    chapter_sections: dict[int, list[str]] = {}
    chapter_start_pages: dict[int, int] = {}
    index = 0
    while index < len(outline):
        item = outline[index]
        title = item.get("/Title", "") if isinstance(item, dict) else ""
        match = re.match(r"^Chapter\s+(\d+)\.\s+", title)
        if match:
            number = int(match.group(1))
            chapter_start_pages[number] = reader.get_destination_page_number(item)
            nested = outline[index + 1] if index + 1 < len(outline) and isinstance(outline[index + 1], list) else []
            chapter_sections[number] = [
                child.get("/Title", "")
                for child in nested
                if isinstance(child, dict) and child.get("/Title")
            ]
        index += 1
    if sorted(chapter_sections) != list(range(1, 29)):
        raise RuntimeError(f"Expected PDF outline sections for political chapters 1–28; found {sorted(chapter_sections)}")
    for record in research["political"]:
        number = record["number"]
        record["sections"] = chapter_sections[number]
        start_page = chapter_start_pages[number]
        next_page = chapter_start_pages.get(number + 1, start_page + 3)
        end_page = min(next_page, start_page + 3)
        raw = "\n".join(reader.pages[page].extract_text() or "" for page in range(start_page, end_page))
        raw = re.sub(rf"^.*?Chapter\s+{number}\.\s+[^\n]+", "", raw, count=1, flags=re.DOTALL | re.IGNORECASE)
        record["summary"] = clean(raw)[:2400]
    research_path.write_text(json.dumps(research, ensure_ascii=False, indent=2), encoding="utf-8")
    return chapter_sections


if __name__ == "__main__":
    ideas = build_ideas()
    history = enrich_history()
    political = enrich_political_history_sections()
    print(json.dumps({"ideas": len(ideas), "availableIdeas": sum(item["status"].startswith("Available") for item in ideas), "plannedIdeas": sum(item["status"] == "Planned" for item in ideas), "suppliedHistoryChapters": len(history), "politicalHistoryChapters": len(political)}, indent=2))
